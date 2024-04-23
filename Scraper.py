import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt  # Points unit

# Create a new Document
doc = Document()

# Access the styles of the document
styles = doc.styles

# Get the Heading 1 style
heading_style = styles['Heading 1']

# Change the font size of the Heading 1 style
heading_style.font.size = Pt(36)  # Sets the font size to 36pt

# # Add a heading of level 1
# doc.add_heading('Chapter 1: Introduction', level=1)

# # Save the document
# doc.save('modified_heading.docx')

# # Initialize a Word document
# doc = Document()

# URL of the table of contents
url = 'https://practicalguidetoevil.wordpress.com/table-of-contents/'

# Send a GET request
response = requests.get(url)
response.raise_for_status()  # Raises an HTTPError for bad responses

# Parse the HTML content
soup = BeautifulSoup(response.text, 'html.parser')

# Find the 'h2' tag with "Book 1"
book_1_header = soup.find('h2', text="Book 1")

# Find all hyperlinks within the next 'ul' sibling of the 'h2' tag
if book_1_header:
    book_1_links = book_1_header.find_next('ul').find_all('a')
    links = [link['href'] for link in book_1_links if link.get('href')]
    # Print or process the links
    for link in links:
        print(link)
else:
    print("Book 1 section not found.")


# List of URLs to scrape
urls = [
    'https://practicalguidetoevil.wordpress.com/2015/04/01/chapter-1-knife/',
    # Add more URLs here
]

for url in links:
    # Fetch the web page
    response = requests.get(url)
    soup = BeautifulSoup(response.content, 'html.parser')

    # Extract the title from the header
    title = soup.find('h1', class_='entry-title').get_text(strip=True)
    doc.add_heading(title, level=1)

    # Extract the content from the entry-content div
    content_div = soup.find('div', class_='entry-content')
    paragraphs = content_div.find_all('p')
    
    first_paragraph = True  # Flag to track the first paragraph
    for p in paragraphs:
        text = p.get_text(strip=True)
        if first_paragraph:  # Check if it's the first paragraph
            p = doc.add_paragraph()
            p.add_run(text).italic = True
            first_paragraph = False  # Update the flag
        else:
            doc.add_paragraph(text)

    # Add a page break after each chapter
    doc.add_page_break()

# Save the document
doc.save('Novel.docx')
